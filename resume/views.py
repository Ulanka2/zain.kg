from rest_framework.views  import  APIView
from rest_framework.response import Response
from rest_framework import status
from resume.models import  PersonalData, Language, Education, WorkExperience, Skills
from resume.serializers import  PersonalDataSerializer, LanguageSerializer
from resume.serializers import EducationSerializer, WorkExperienceSerializer, SkillsSerializer
from rest_framework.generics import get_object_or_404
from docx import Document
from docx.shared import Inches


class PersonalDataAPIView(APIView):

    def post(self, request, format=None):
        serializer = PersonalDataSerializer(data=request.data)
        if serializer.is_valid():
            instance = serializer.save()
            doc = Document('media/example.docx')
            tbl = doc.tables[0]
            cell = tbl.rows[0].cells[0]
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run()
            run.add_picture(instance.image, height=Inches(1.90), width=Inches(1.80))
            table = doc.tables[1]
            hdr_cells = table.rows[0].cells
            hdr_cells[1].text = instance.position
            hdr_cells = table.rows[1].cells
            hdr_cells[1].text = instance.num_passport
            table = doc.tables[2]
            hdr_cells = table.rows[1].cells
            hdr_cells[1].text = instance.full_name
            hdr_cells = table.rows[2].cells
            hdr_cells[1].text = instance.nationality_gender
            hdr_cells = table.rows[3].cells
            hdr_cells[1].text = instance.country_city_of_residence
            hdr_cells = table.rows[4].cells
            hdr_cells[1].text = instance.date_of_birth
            hdr_cells = table.rows[5].cells
            hdr_cells[1].text = instance.age_height_weight
            hdr_cells = table.rows[6].cells
            hdr_cells[1].text = instance.status_children
            hdr_cells = table.rows[7].cells
            hdr_cells[1].text = instance.health_smoker
            doc.add_picture(instance.image_full_height, width=Inches(5))
            name = instance.full_name.replace(' ', '')
            doc_name = f'media/media_rezume/{name}-{instance.id}.docx'
            doc.save(doc_name)
            doc_name =  doc_name[6:]
            instance.file = doc_name
            instance.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class LanguageAPIView(APIView):

    def post(self, request, id):
        personal_data = get_object_or_404(PersonalData, id=id)
        serializer = LanguageSerializer(data=request.data)
        if serializer.is_valid():
            if personal_data.language_set.all().order_by('-id').count() >= 4:
                return Response({'error: You have too many languages'}, status=status.HTTP_400_BAD_REQUEST)
            
            instance = serializer.save(personal_data=personal_data)
            name = personal_data.full_name.replace(' ', '')
            doc = Document(f'media/media_rezume/{name}-{personal_data.id}.docx')
            table = doc.tables[3]
            personal_data = instance.personal_data
            languages = personal_data.language_set.all()[:4]
            basic_number = 2
            for instance in languages:
                hdr_cells = table.rows[basic_number].cells
                hdr_cells[0].text = instance.language.name
                table = doc.tables[3]
                hdr_cells = table.rows[basic_number].cells
                hdr_cells[int(instance.written) -1].text = instance.written
                hdr_cells[int(instance.spoken) +3].text = instance.spoken
                hdr_cells[int(instance.understanding) +7].text = instance.understanding
                basic_number += 1
            name = personal_data.full_name.replace(' ', '')
            doc_name = f'media/media_rezume/{name}-{personal_data.id}.docx'
            doc.save(doc_name)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response({'detail': serializer.errors}, status=status.HTTP_400_BAD_REQUEST)


class EducationAPIView(APIView):

    def post(self, request, id):
        personal_data = get_object_or_404(PersonalData, id=id)
        serializer = EducationSerializer(data=request.data)
        if serializer.is_valid():
            if personal_data.education_set.all().order_by('-id').count() >= 3:
                return Response({'error: You have too many education'}, status=status.HTTP_400_BAD_REQUEST)
            instance = serializer.save(personal_data=personal_data)
            name = personal_data.full_name.replace(' ', '')
            doc = Document(f'media/media_rezume/{name}-{personal_data.id}.docx')
            table = doc.tables[4]
            personal_data = instance.personal_data
            education = personal_data.education_set.all()[:3]
            basic_number = 2
            for instance in education:
                hdr_cells = table.rows[basic_number].cells
                hdr_cells[0].text = instance.university
                hdr_cells = table.rows[basic_number].cells
                hdr_cells[1].text = instance.specialization
                hdr_cells = table.rows[basic_number].cells
                hdr_cells[2].text = instance.duration
                hdr_cells = table.rows[basic_number].cells
                hdr_cells[3].text = instance.city_country
                basic_number += 1 
            name = personal_data.full_name.replace(' ', '')
            doc_name = f'media/media_rezume/{name}-{personal_data.id}.docx'
            doc.save(doc_name)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response({'detail': serializer.errors}, status=status.HTTP_400_BAD_REQUEST)


class WorkExperienceOneAPIView(APIView):

    def post(self, request, id):
        personal_data = get_object_or_404(PersonalData, id=id)
        serializer = WorkExperienceSerializer(data=request.data)
        if serializer.is_valid():
            instance = serializer.save(personal_data=personal_data)
            name = personal_data.full_name.replace(' ', '')
            doc = Document(f'media/media_rezume/{name}-{personal_data.id}.docx')
            table = doc.tables[5]
            hdr_cells = table.rows[1].cells
            hdr_cells[1].text = instance.position
            hdr_cells = table.rows[2].cells
            hdr_cells[1].text = instance.company_name
            hdr_cells = table.rows[3].cells
            hdr_cells[1].text = instance.period_city_country
            hdr_cells = table.rows[4].cells
            hdr_cells[1].text = instance.responsibilities
            name = personal_data.full_name.replace(' ', '')
            doc_name = f'media/media_rezume/{name}-{personal_data.id}.docx'
            doc.save(doc_name)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response({'detail': serializer.errors}, status=status.HTTP_400_BAD_REQUEST)


class WorkExperienceTwoAPIView(APIView):

    def post(self, request, id):
        personal_data = get_object_or_404(PersonalData, id=id)
        serializer = WorkExperienceSerializer(data=request.data)
        if serializer.is_valid():
            instance = serializer.save(personal_data=personal_data)
            name = personal_data.full_name.replace(' ', '')
            doc = Document(f'media/media_rezume/{name}-{personal_data.id}.docx')
            table = doc.tables[5]
            hdr_cells = table.rows[6].cells
            hdr_cells[1].text = instance.position
            hdr_cells = table.rows[7].cells
            hdr_cells[1].text = instance.company_name
            hdr_cells = table.rows[8].cells
            hdr_cells[1].text = instance.period_city_country
            hdr_cells = table.rows[9].cells
            hdr_cells[1].text = instance.responsibilities
            name = personal_data.full_name.replace(' ', '')
            doc_name = f'media/media_rezume/{name}-{personal_data.id}.docx'
            doc.save(doc_name)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response({'detail': serializer.errors}, status=status.HTTP_400_BAD_REQUEST)


class WorkExperienceThreeAPIView(APIView):

    def post(self, request, id):
        personal_data = get_object_or_404(PersonalData, id=id)
        serializer = WorkExperienceSerializer(data=request.data)
        if serializer.is_valid():
            instance = serializer.save(personal_data=personal_data)
            name = personal_data.full_name.replace(' ', '')
            doc = Document(f'media/media_rezume/{name}-{personal_data.id}.docx')
            table = doc.tables[5]
            hdr_cells = table.rows[11].cells
            hdr_cells[1].text = instance.position
            hdr_cells = table.rows[12].cells
            hdr_cells[1].text = instance.company_name
            hdr_cells = table.rows[13].cells
            hdr_cells[1].text = instance.period_city_country
            hdr_cells = table.rows[14].cells
            hdr_cells[1].text = instance.responsibilities
            name = personal_data.full_name.replace(' ', '')
            doc_name = f'media/media_rezume/{name}-{personal_data.id}.docx'
            doc.save(doc_name)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response({'detail': serializer.errors}, status=status.HTTP_400_BAD_REQUEST)


class SkillsAPIView(APIView):

    def post(self, request, id):
        personal_data = get_object_or_404(PersonalData, id=id)
        serializer = SkillsSerializer(data=request.data)
        
        if serializer.is_valid():
            instance = serializer.save(personal_data=personal_data)
            name = personal_data.full_name.replace(' ', '')
            doc = Document(f'media/media_rezume/{name}-{personal_data.id}.docx')
            table = doc.tables[6]
            hdr_cells = table.rows[1].cells
            hdr_cells[0].text = instance.skills
            
            table = doc.tables[7]
            if instance.work_time  == 'Less than 40 hours/week':
                hdr_cells = table.rows[1].cells
            elif instance.work_time  == '40–50 hours/week':
                hdr_cells = table.rows[2].cells
            else:
                hdr_cells = table.rows[3].cells
            hdr_cells[2].text = 'True'

            table = doc.tables[8]
            hdr_cells = table.rows[0].cells
            hdr_cells[1].text = instance.i_confirm
            name = personal_data.full_name.replace(' ', '')
            doc_name = f'media/media_rezume/{name}-{personal_data.id}.docx'
            doc.save(doc_name)
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        return Response({'detail': serializer.errors}, status=status.HTTP_400_BAD_REQUEST)
